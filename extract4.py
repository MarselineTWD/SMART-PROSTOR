# -*- coding: utf-8 -*-
import sys, io, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
cands = glob.glob(r"Файлы\**\*2.1*.docx", recursive=True)
print("FOUND:", cands)
if cands:
    doc=Document(cands[0])
    print("--- PARAGRAPHS (non-empty) ---")
    n=0
    for p in doc.paragraphs:
        t=p.text.strip()
        if t:
            print(t[:140]); n+=1
        if n>80: break
    for i,t in enumerate(doc.tables):
        print(f"\n--- TABLE {i} ({len(t.rows)}x{len(t.columns)}) ---")
        for r in t.rows[:12]:
            cells=[c.text.strip().replace('\n',' ')[:60] for c in r.cells]
            uniq=[]
            for c in cells:
                if not uniq or uniq[-1]!=c: uniq.append(c)
            print(" | ".join(uniq)[:190])
