# -*- coding: utf-8 -*-
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

targets = [
 r"Файлы\Выгрузка из системы\Шаблоны ТЗ\ТЗ Концепт геологии.docx",
 r"Файлы\Приложение № 2.1 Форма Технического задания.docx",
]
for path in targets:
    print("\n"+"#"*80); print("FILE:", os.path.basename(path)); print("#"*80)
    doc=Document(path)
    print("\n--- PARAGRAPHS ---")
    for p in doc.paragraphs:
        t=p.text.strip()
        if t: print(t[:150])
    for i,t in enumerate(doc.tables):
        print(f"\n--- TABLE {i} ({len(t.rows)}x{len(t.columns)}) ---")
        for r in t.rows:
            cells=[c.text.strip().replace("\n"," ")[:70] for c in r.cells]
            # collapse duplicates from merged cells
            uniq=[]
            for c in cells:
                if not uniq or uniq[-1]!=c: uniq.append(c)
            print(" | ".join(uniq)[:200])
