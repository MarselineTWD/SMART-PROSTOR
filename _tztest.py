# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from backend.app.models.domain import DraftInputData
from backend.app.services.tz_templates import tz_template_service
from backend.app.services.tz_generator import tz_generator
from backend.app.services.documents import document_export_service

print("templates:", len(tz_template_service.list_templates()))
tpl = tz_template_service.get_template("tz-ptd-reserves")
print("template:", tpl.name, "| sections:", len(tpl.sections))

doc = tz_generator.new_document(
    tpl,
    object_name="Приразломное месторождение",
    customer_name="АО «Газпромнефть-ННГ»",
    executor_name="ООО «НТЦ»",
    input_data=DraftInputData(goal="выполнить подсчёт запасов и подготовить ПТД",
                              deadline="25.12.2026", source_data_ready=True,
                              requires_subcontractor=True, subcontract_share_percent=40,
                              separate_subcontract_estimate=True),
    requisites={"city": "г. Тюмень", "contract_number": "123/26", "contract_date": "12.01.2026"},
)
print("empty ready_score:", doc.ready_score)

# augment (should fill all empty)
tz_generator.generate(doc, mode="augment")
print("after augment ready_score:", doc.ready_score)
print("filled sections:", sum(1 for s in doc.sections if s.content.strip()), "/", len(doc.sections))

print("\n=== ЦЕЛИ ===\n", doc.sections[0].content)
sub = next(s for s in doc.sections if s.key=="subcontractors")
print("\n=== СУБПОДРЯД ===\n", sub.content)

path = document_export_service.export_docx(doc)
print("\nDOCX:", path.name, "exists:", path.exists(), "bytes:", path.stat().st_size)

# validate docx reopen
from docx import Document
d2 = Document(str(path))
print("docx paragraphs:", len(d2.paragraphs), "tables:", len(d2.tables))

# legacy export from draft
from backend.app.schemas.draft import DraftFromSearchRequest
from backend.app.services.drafts import draft_service
draft = draft_service.create_from_search(DraftFromSearchRequest(product_id="product-reserves",
        input_data=DraftInputData(object_name="Северный блок", customer_name="Блок геологии",
                                  goal="оценка запасов", deadline="2026-09-30")))
p2 = document_export_service.export_from_draft(draft)
print("legacy docx:", p2.name, "exists:", p2.exists())
print("ALL_OK")
