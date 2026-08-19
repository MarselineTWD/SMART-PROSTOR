"""Экспорт технического задания в DOCX.

Формируется единый документ .docx на основе сохранённого ТЗ (``TZDocument``).
Ранее пакет включал XLSX (КП/РС) — от них отказались: результат один
самодостаточный файл ТЗ.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from backend.app.models.domain import RequestDraft, TZDocument
from backend.app.services.tz_generator import tz_generator
from backend.app.services.tz_templates import tz_template_service


PROJECT_ROOT = Path(__file__).resolve().parents[3]
OUTPUT_ROOT = PROJECT_ROOT / "outputs"
TEMPLATE_ROOTS = (
    Path("/app/templates"),
    PROJECT_ROOT / "Файлы" / "Выгрузка из системы" / "Шаблоны ТЗ",
)


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-zА-Яа-яЁё _.-]+", "", name).strip()
    return (cleaned or "ТЗ")[:80]


def _plain_ai_text(value: object, *, strip: bool = True) -> str:
    """Убирает служебную markdown-разметку, не предназначенную для DOCX."""
    text = str(value or "")
    text = re.sub(r"```(?:[a-zA-Z0-9_-]+)?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"(?<!\\)[*_]{1,3}", "", text)
    text = text.replace("`", "")
    text = re.sub(r"(?<!\w)#{1,6}\s*", "", text)
    return text.strip() if strip else text


class DocumentExportService:
    """Строит DOCX-файл технического задания."""

    def export_docx(self, document: TZDocument) -> Path:
        package_dir = OUTPUT_ROOT / document.id
        package_dir.mkdir(parents=True, exist_ok=True)
        path = package_dir / f"{_safe_filename(document.title or 'Техническое задание')}.docx"
        self._build_tz(document, path)
        return path

    def export_from_draft(self, draft: RequestDraft) -> Path:
        """Совместимость: собирает ТЗ из RequestDraft и выгружает DOCX."""
        document = tz_generator.document_from_draft(draft)
        return self.export_docx(document)

    def _build_tz(self, document: TZDocument, path: Path) -> None:
        doc, source_name = self._document_from_source_template(document)
        if source_name is None:
            normal = doc.styles["Normal"]
            normal.font.name = "Arial"
            normal.font.size = Pt(11)

        title = self._add_heading(doc, "Техническое задание", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if document.template_name:
            subtitle = doc.add_paragraph(_plain_ai_text(document.template_name))
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if subtitle.runs:
                subtitle.runs[0].bold = True
                subtitle.runs[0].font.size = Pt(14)
        stamp = doc.add_paragraph(f"Сформировано в ПРОСТОР 2.0: {datetime.now():%d.%m.%Y %H:%M}")
        stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if source_name:
            source = doc.add_paragraph(f"Оформление: исходный шаблон «{source_name}»")
            source.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if source.runs:
                source.runs[0].italic = True
                source.runs[0].font.color.rgb = RGBColor(90, 100, 115)

        self._add_key_value_table(
            doc,
            [
                ("Договор", self._contract_line(document)),
                ("Объект", document.object_name or "Не заполнено"),
                ("Заказчик", document.customer_name or "Не заполнено"),
                ("Исполнитель", document.executor_name or "Не выбран"),
                ("Место выполнения", str(document.requisites.get("city") or "Не заполнено")),
                ("Готовность", f"{document.ready_score}%"),
            ],
        )

        services = document.requisites.get("services") or []
        if services:
            self._add_heading(doc, "Состав услуг", level=1)
            for item in services:
                service = item if isinstance(item, dict) else {"name": str(item)}
                paragraph = doc.add_paragraph(style=self._style_or_none(doc, "List Bullet"))
                run = paragraph.add_run(_plain_ai_text(service.get("name")))
                run.bold = bool(service.get("mandatory"))
                if service.get("mandatory"):
                    paragraph.add_run(" — обязательная услуга")

        stages = [str(item) for item in (document.requisites.get("stages") or []) if str(item).strip()]
        if stages:
            self._add_heading(doc, "Этапы и календарный план", level=1)
            plan_table = doc.add_table(rows=1, cols=2)
            plan_table.style = self._table_style(doc)
            plan_table.rows[0].cells[0].text = "№"
            plan_table.rows[0].cells[1].text = "Этап работ"
            for index, stage in enumerate(stages, start=1):
                cells = plan_table.add_row().cells
                cells[0].text = str(index)
                cells[1].text = _plain_ai_text(stage)
            constraints = document.requisites.get("schedule_constraints") or []
            for constraint in constraints:
                after = int(constraint.get("after_stage") or 0)
                reason = _plain_ai_text(constraint.get("reason") or "Приостановка работ")
                days = int(constraint.get("days") or 0)
                doc.add_paragraph(
                    f"Пауза после этапа {after or 'до начала работ'}: {reason}, {days} календ. дн.",
                    style=self._style_or_none(doc, "Intense Quote"),
                )

        for index, section in enumerate(document.sections, start=1):
            self._add_heading(doc, f"{index}. {section.title}", level=1)
            content = section.content.strip() or "Раздел требует заполнения."
            self._add_formatted_text(doc, content)

        self._add_heading(doc, "Подписи сторон", level=1)
        self._add_key_value_table(
            doc,
            [
                ("ЗАКАЗЧИК", str(document.requisites.get("signatory_customer") or "____________ / ____________")),
                ("ИСПОЛНИТЕЛЬ", str(document.requisites.get("signatory_executor") or "____________ / ____________")),
            ],
        )

        doc.save(path)

    def _document_from_source_template(self, document: TZDocument):
        """Открывает реальный DOCX как носитель стилей, колонтитулов и разметки страницы."""
        template = tz_template_service.get_template(document.template_key)
        for filename in (template.source_files if template else []):
            for root in TEMPLATE_ROOTS:
                if not root.exists():
                    continue
                candidates = [root / filename, *root.rglob(filename)]
                for candidate in candidates:
                    if not candidate.is_file():
                        continue
                    try:
                        doc = Document(str(candidate))
                        body = doc._element.body
                        for child in list(body):
                            if child.tag != qn("w:sectPr"):
                                body.remove(child)
                        return doc, filename
                    except (OSError, ValueError, KeyError):
                        continue
        return Document(), None

    @staticmethod
    def _style_or_none(doc: Document, name: str):
        return next(
            (style for style in doc.styles if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name == name),
            None,
        )

    def _add_heading(self, doc: Document, text: str, *, level: int):
        preferred = "Title" if level == 0 else f"Heading {level}"
        style = self._style_or_none(doc, preferred)
        paragraph = doc.add_paragraph(_plain_ai_text(text), style=style)
        if style is None and paragraph.runs:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = Pt(18 if level == 0 else 14)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    @staticmethod
    def _table_style(doc: Document) -> str | None:
        return next(
            (style for style in doc.styles if style.type == WD_STYLE_TYPE.TABLE and style.name == "Table Grid"),
            None,
        )

    def _add_formatted_text(self, doc: Document, content: str) -> None:
        """Сохраняет смысловую структуру текста: списки, абзацы и выделения."""
        for raw in content.splitlines():
            line = raw.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("```"):
                continue
            markdown_heading = re.match(r"^#{1,6}\s*(.+)$", line)
            if markdown_heading:
                heading = _plain_ai_text(markdown_heading.group(1))
                if heading:
                    self._add_heading(doc, heading, level=2)
                continue
            numbered = re.match(r"^\d+[.)]\s+(.*)$", line)
            bulleted = re.match(r"^[•*\-–—]\s+(.*)$", line)
            if numbered:
                paragraph = doc.add_paragraph(style=self._style_or_none(doc, "List Number"))
                self._add_inline_text(paragraph, numbered.group(1))
            elif bulleted:
                paragraph = doc.add_paragraph(style=self._style_or_none(doc, "List Bullet"))
                self._add_inline_text(paragraph, bulleted.group(1))
            else:
                paragraph = doc.add_paragraph()
                self._add_inline_text(paragraph, line)

    @staticmethod
    def _add_inline_text(paragraph, value: str) -> None:
        """Конвертирует **выделение** в DOCX-run и удаляет прочий markdown."""
        cursor = 0
        for match in re.finditer(r"\*\*(.+?)\*\*|__(.+?)__", value):
            if match.start() > cursor:
                paragraph.add_run(_plain_ai_text(value[cursor:match.start()], strip=False))
            run = paragraph.add_run(_plain_ai_text(match.group(1) or match.group(2)))
            run.bold = True
            cursor = match.end()
        if cursor < len(value):
            paragraph.add_run(_plain_ai_text(value[cursor:], strip=False))

    @staticmethod
    def _contract_line(document: TZDocument) -> str:
        number = document.requisites.get("contract_number")
        date = document.requisites.get("contract_date")
        if number or date:
            return f"№ {number or '—'} от {date or '—'}"
        return document.contract_name or "Не выбран"

    @staticmethod
    def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
        table = doc.add_table(rows=0, cols=2)
        table.style = DocumentExportService._table_style(doc)
        for key, value in rows:
            cells = table.add_row().cells
            cells[0].text = _plain_ai_text(key)
            cells[1].text = _plain_ai_text(value)


document_export_service = DocumentExportService()
