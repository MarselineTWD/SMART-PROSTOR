# -*- coding: utf-8 -*-
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

base = r"Файлы\Выгрузка из системы\Шаблоны ТЗ"
extra = [r"Файлы\Приложение № 2.1 Форма Технического задания.docx",
         r"Файлы\ТЗ_Хакатон_ПРОСТОР (финал)-1.docx"]
files=[]
for root,_,fs in os.walk(base):
    for f in fs:
        if f.lower().endswith('.docx'): files.append(os.path.join(root,f))
files += [e for e in extra if os.path.exists(e)]

sec_re = re.compile(r'^\s*(\d{1,2})\.\s+(\S.+)')
ph_re = re.compile(r'\{[^}]{2,60}\}')

for path in files:
    print("\n"+"#"*80); print("FILE:", os.path.basename(path))
    try: doc=Document(path)
    except Exception as e: print("  err",e); continue
    def alltext():
        for p in doc.paragraphs: yield p.text
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: yield c.text
    secs=[]
    for p in doc.paragraphs:
        m=sec_re.match(p.text.strip())
        if m and len(p.text.strip())<120:
            secs.append(f"{m.group(1)}. {m.group(2)[:80]}")
    print(" SECTIONS:")
    for s in secs[:20]: print("   -",s)
    phs=set()
    for txt in alltext():
        for m in ph_re.findall(txt): phs.add(m)
    if phs:
        print(" PLACEHOLDERS:", " | ".join(sorted(phs))[:600])
