import asyncio
import shutil
import unittest
from pathlib import Path
from unittest.mock import AsyncMock, patch

from docx import Document as DocxDocument

from backend.app.models.domain import DraftInputData, TZDocumentSection, TZFeedback
from backend.app.api.routes.tz import complete_document
from backend.app.schemas.draft import DraftFromSearchRequest
from backend.app.schemas.tz import TZCompleteRequest
from backend.app.services.documents import document_export_service
from backend.app.services.drafts import draft_service
from backend.app.services.procurement import procurement_service
from backend.app.services.search import search_service
from backend.app.services.tz_generator import tz_generator
from backend.app.services.tz_templates import tz_template_service
from backend.app.services.tz_validation import tz_validation_service


class ProstorMvpTest(unittest.TestCase):
    def test_docx_export_converts_ai_markdown_to_native_formatting(self):
        doc = DocxDocument()
        document_export_service._add_formatted_text(
            doc,
            "### Результаты\n**Важно:** проверить данные\n- `Итоговый` отчёт\n```text\nБез ограждения\n```",
        )
        document_export_service._add_heading(doc, "1. ### Служебный заголовок", level=1)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        for marker in ("###", "**", "```", "`"):
            self.assertNotIn(marker, text)
        self.assertIn("Результаты", text)
        self.assertIn("Важно: проверить данные", text)
        self.assertTrue(any(run.bold and "Важно:" in run.text for p in doc.paragraphs for run in p.runs))

    def test_reserves_query_returns_reserves_product(self):
        result = search_service.search(
            "Нужно оценить запасы по объекту и подготовить проектно-технический документ",
            limit=3,
        )

        self.assertGreaterEqual(len(result.products), 1)
        self.assertEqual(result.products[0].product.id, "product-reserves")
        self.assertGreater(len(result.products[0].recommended_companies), 0)
        self.assertGreater(len(result.products[0].similar_cases), 0)

    def test_3d_without_input_data_creates_risks(self):
        draft = draft_service.create_from_search(
            DraftFromSearchRequest(
                product_id="product-reserves",
                input_data=DraftInputData(
                    customer_name="Блок геологии и разработки",
                    goal="Оценить запасы и подготовить ПТД",
                    deadline="2026-09-30",
                    needs_3d_model=True,
                    source_data_ready=False,
                ),
            )
        )

        risk_codes = {risk.code for risk in draft.risks}
        self.assertIn("missing_object", risk_codes)
        self.assertIn("3d_missing_input_data", risk_codes)
        self.assertLess(draft.ready_score, 100)

    def test_export_produces_single_docx_without_xlsx(self):
        draft = draft_service.create_from_search(
            DraftFromSearchRequest(
                product_id="product-reserves",
                input_data=DraftInputData(
                    object_name="Северный блок",
                    customer_name="Блок геологии и разработки",
                    goal="Оценить запасы и подготовить ПТД",
                    deadline="2026-09-30",
                ),
            )
        )

        docx_path = document_export_service.export_from_draft(draft)
        try:
            self.assertTrue(docx_path.exists())
            self.assertEqual(docx_path.suffix, ".docx")
            # Никаких xlsx рядом с результатом.
            self.assertEqual(list(docx_path.parent.glob("*.xlsx")), [])
            # Файл открывается как валидный docx.
            exported = DocxDocument(str(docx_path))
            self.assertGreater(len(exported.paragraphs), 5)
            exported_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
            self.assertIn("Оформление: исходный шаблон", exported_text)
            self.assertIn("Состав услуг", exported_text)
            self.assertIn("Этапы и календарный план", exported_text)
            self.assertGreaterEqual(len(exported.tables), 3)
        finally:
            shutil.rmtree(Path(docx_path).parent, ignore_errors=True)


class TZTemplatesTest(unittest.TestCase):
    def test_services_and_plan_depend_on_tz_conditions(self):
        template = tz_template_service.get_template("tz-geology-concept")
        simple = tz_generator.new_document(
            template,
            input_data=DraftInputData(source_data_ready=True),
        )
        complex_tz = tz_generator.new_document(
            template,
            input_data=DraftInputData(
                source_data_ready=False,
                needs_3d_model=True,
                requires_subcontractor=True,
                separate_subcontract_estimate=True,
            ),
        )
        complex_tz.requisites["services"].append({
            "name": "Петрофизическая интерпретация",
            "mandatory": False,
            "source": "manual",
        })

        tz_generator.generate(simple, plan_only=True, template=template)
        tz_generator.generate(complex_tz, plan_only=True, template=template)

        service_names = {item["name"] for item in complex_tz.requisites["services"]}
        self.assertIn("Построение и проверка 3D-модели", service_names)
        self.assertIn("Петрофизическая интерпретация", service_names)
        self.assertNotEqual(simple.requisites["stages"], complex_tz.requisites["stages"])
        self.assertGreater(len(complex_tz.requisites["stages"]), len(simple.requisites["stages"]))

        removed = next(item for item in complex_tz.requisites["services"] if item["source"] == "rule")
        complex_tz.requisites["services"].remove(removed)
        complex_tz.requisites["removed_auto_services"] = [removed["name"]]
        tz_generator.generate(complex_tz, plan_only=True, template=template)
        self.assertNotIn(removed["name"], {item["name"] for item in complex_tz.requisites["services"]})

    def test_catalog_has_all_templates(self):
        templates = tz_template_service.list_templates()
        self.assertEqual(len(templates), 11)
        for tpl in templates:
            self.assertTrue(tpl.sections, f"{tpl.key} без разделов")
            self.assertTrue(tpl.example.get("stages"), f"{tpl.key} без примера")
        self.assertIsNotNone(tz_template_service.get_template("tz-ptd-opz"))
        self.assertIsNone(tz_template_service.get_template("does-not-exist"))

    def test_template_types_have_distinct_dynamic_fields(self):
        geology = tz_template_service.get_template("tz-geology-concept")
        completion = tz_template_service.get_template("tz-integrated-completion")
        geology_keys = {field.key for field in geology.fields}
        completion_keys = {field.key for field in completion.fields}
        self.assertIn("target_horizons", geology_keys)
        self.assertNotIn("target_horizons", completion_keys)
        self.assertIn("completion_type", completion_keys)
        self.assertEqual(
            next(field for field in completion.fields if field.key == "completion_type").input_type,
            "select",
        )

    def test_template_for_product_falls_back_to_universal(self):
        self.assertEqual(
            tz_template_service.template_for_product("product-geology").key,
            "tz-geology-concept",
        )
        self.assertEqual(
            tz_template_service.template_for_product("unknown-product").key,
            "tz-universal",
        )


class TZGeneratorTest(unittest.TestCase):
    def _doc(self):
        tpl = tz_template_service.get_template("tz-ptd-reserves")
        return tpl, tz_generator.new_document(
            tpl,
            object_name="Приразломное месторождение",
            customer_name="АО «Заказчик»",
            input_data=DraftInputData(goal="подсчёт запасов", deadline="25.12.2026"),
        )

    def test_full_generation_fills_all_sections(self):
        tpl, doc = self._doc()
        tz_generator.generate(doc, mode="full", template=tpl)
        self.assertTrue(all(s.content.strip() for s in doc.sections))
        self.assertTrue(all(s.source == "ai" for s in doc.sections))
        self.assertGreaterEqual(doc.ready_score, 60)
        self.assertTrue(doc.ai_initially_generated)

    def test_custom_section_is_generated_using_its_title(self):
        tpl, doc = self._doc()
        doc.sections.append(TZDocumentSection(
            key="custom-security", title="Требования к информационной безопасности"
        ))
        tz_generator.generate(doc, section_keys=["custom-security"], template=tpl)
        custom = doc.sections[-1]
        self.assertIn("Требования к информационной безопасности", custom.content)
        self.assertEqual(custom.source, "ai")

    def test_feedback_rating_is_limited_to_five_point_scale(self):
        feedback = TZFeedback(
            contractor={"rating": 5, "comment": "Работы выполнены качественно"},
            ai_tz={"rating": 4, "comment": "Потребовалась небольшая доработка"},
            ai_chat={"rating": 5, "comment": "Полезные ответы"},
        )
        self.assertEqual(feedback.contractor.rating, 5)
        with self.assertRaises(ValueError):
            TZFeedback(contractor={"rating": 6})

    def test_augment_keeps_manual_sections(self):
        tpl, doc = self._doc()
        doc.sections[0].content = "Ручной текст цели"
        doc.sections[0].source = "manual"
        tz_generator.generate(doc, mode="augment", template=tpl)
        self.assertEqual(doc.sections[0].content, "Ручной текст цели")
        self.assertEqual(doc.sections[0].source, "manual")
        self.assertTrue(doc.sections[1].content.strip())

    def test_section_keys_regenerate_only_selected(self):
        tpl, doc = self._doc()
        tz_generator.generate(doc, mode="full", template=tpl)
        snapshot = {s.key: s.content for s in doc.sections}
        tz_generator.generate(
            doc,
            mode="augment",
            section_keys=["goal"],
            instruction="учесть сжатые сроки",
            template=tpl,
        )
        for section in doc.sections:
            if section.key == "goal":
                self.assertNotEqual(section.content, snapshot["goal"])
                self.assertIn("сжатые сроки", section.content)
            else:
                self.assertEqual(section.content, snapshot[section.key])


class ProcurementEstimateTest(unittest.TestCase):
    def test_products_loaded_from_seed(self):
        products = procurement_service.list_products()
        self.assertGreaterEqual(len(products), 20)
        sample = products[0]
        for key in ("product_id", "name", "company_count", "min_days", "max_days"):
            self.assertIn(key, sample)
        self.assertGreater(sample["company_count"], 0)

    def test_match_products_by_name(self):
        matches = procurement_service.match_products("ТЗ: Концепт геологии", limit=3)
        self.assertTrue(matches)
        self.assertEqual(matches[0]["name"], "Концепт геологии")

    def test_estimate_sorted_and_roadmap_consistent(self):
        product_id = procurement_service.list_products()[0]["product_id"]
        estimate = procurement_service.estimate_product(product_id)
        self.assertIsNotNone(estimate)
        companies = estimate["companies"]
        self.assertGreater(len(companies), 0)
        days = [c["estimated_days"] for c in companies]
        self.assertEqual(days, sorted(days))  # быстрейший подрядчик — первым
        for company in companies:
            stages = company["stages"]
            if not stages:
                continue
            self.assertEqual(sum(s["days"] for s in stages), company["estimated_days"])
            offset = 0
            for stage in stages:
                self.assertEqual(stage["offset_days"], offset)
                offset += stage["days"]

    def test_estimate_missing_product_returns_none(self):
        self.assertIsNone(procurement_service.estimate_product("no-such-product"))

    def test_estimate_contains_explainable_cost_for_every_contractor(self):
        product = procurement_service.list_products()[0]
        estimate = procurement_service.estimate_product(product["product_id"])
        self.assertIsNotNone(estimate)
        self.assertGreater(estimate["summary"]["lowest_cost_without_vat"], 0)
        self.assertIn("L2", estimate["summary"]["cost_disclaimer"])
        self.assertIn("12 000", estimate["summary"]["cost_disclaimer"])
        for company in estimate["companies"]:
            self.assertGreater(company["cost_without_vat"], 0)
            self.assertEqual(company["cost_confidence"], "indicative")
            self.assertAlmostEqual(
                company["cost_with_vat"],
                company["cost_without_vat"] + company["vat_amount"],
            )

    def test_additional_service_recalculates_matching_contractors(self):
        base = None
        for product in procurement_service.list_products():
            candidate = procurement_service.estimate_product(product["product_id"])
            if candidate and candidate["available_additional_services"]:
                base = candidate
                break
        self.assertIsNotNone(base)
        addon_id = base["available_additional_services"][0]["product_id"]
        recalculated = procurement_service.estimate_product(
            base["product_id"], additional_product_ids=[addon_id]
        )
        self.assertEqual(recalculated["selected_additional_product_ids"], [addon_id])
        affected = [c for c in recalculated["companies"] if c["additional_cost_without_vat"] > 0]
        self.assertTrue(affected)
        for company in affected:
            self.assertEqual(
                company["cost_without_vat"],
                company["base_cost_without_vat"] + company["additional_cost_without_vat"],
            )

    def test_contractors_without_selected_service_are_excluded(self):
        base = None
        addon = None
        for product in procurement_service.list_products():
            candidate = procurement_service.estimate_product(product["product_id"])
            partial = next((item for item in candidate["available_additional_services"]
                            if item["common_company_count"] < candidate["summary"]["company_count"]), None)
            if partial:
                base, addon = candidate, partial
                break
        self.assertIsNotNone(base)

        result = procurement_service.estimate_product(
            base["product_id"], additional_product_ids=[addon["product_id"]]
        )
        self.assertEqual(len(result["companies"]), addon["common_company_count"])
        self.assertEqual(
            result["summary"]["excluded_company_count"],
            base["summary"]["company_count"] - addon["common_company_count"],
        )
        self.assertTrue(all(company["additional_services"] for company in result["companies"]))
        self.assertTrue(all(row["reasons"] for row in result["excluded_contractors"]))

    def test_impossible_deadline_excludes_contractors_with_reason(self):
        product_id = procurement_service.list_products()[0]["product_id"]
        result = procurement_service.estimate_product(
            product_id,
            project_context={
                "requisites": {"start_date": "2026-08-19"},
                "input_data": {"deadline": "2026-08-20"},
            },
        )
        self.assertEqual(result["companies"], [])
        self.assertGreater(result["summary"]["excluded_company_count"], 0)
        self.assertTrue(all(
            any("Срок недостижим" in reason for reason in company["reasons"])
            for company in result["excluded_contractors"]
        ))

    def test_tz_roadmap_uses_only_stages_written_in_document(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(
            template,
            title="ТЗ: Концепт геологии Северного участка",
            object_name="Северный участок",
            customer_name="Блок геологии",
            input_data=DraftInputData(goal="Подготовить геологическую концепцию", deadline="2026-12-20"),
        )
        expected_stages = ["Проверка исходных данных", "Интерпретация", "Выпуск отчёта"]
        document.requisites["stages"] = expected_stages

        response = procurement_service.estimate_for_tz(document)
        estimate = response["estimate"]

        self.assertIsNotNone(estimate)
        self.assertEqual(estimate["roadmap_source"], "tz")
        self.assertEqual(estimate["tz_id"], document.id)
        for company in estimate["companies"]:
            self.assertEqual([stage["name"] for stage in company["stages"]], expected_stages)
            self.assertEqual(sum(stage["days"] for stage in company["stages"]), company["estimated_days"])

    def test_tz_without_stages_has_no_roadmap(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(template, title="ТЗ: Концепт геологии")
        document.requisites["stages"] = []
        self.assertIsNone(procurement_service.estimate_for_tz(document)["estimate"])

    def test_tz_roadmap_includes_explicit_pause_and_its_cost(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(
            template,
            title="ТЗ: Концепт геологии",
            input_data=DraftInputData(goal="Подготовить концепт геологии"),
        )
        document.requisites["stages"] = ["Сбор данных", "Интерпретация", "Отчёт"]
        document.requisites["schedule_constraints"] = [{
            "after_stage": 1,
            "reason": "Ожидание исходных данных",
            "days": 7,
            "billable_percent": 25,
        }]
        estimate = procurement_service.estimate_for_tz(document)["estimate"]
        self.assertIsNotNone(estimate)
        for company in estimate["companies"]:
            pauses = [stage for stage in company["stages"] if stage.get("kind") == "pause"]
            self.assertEqual(len(pauses), 1)
            self.assertEqual(pauses[0]["days"], 7)
            self.assertEqual(pauses[0]["billable_percent"], 25)
            self.assertGreater(pauses[0]["estimated_cost_without_vat"], 0)
            self.assertEqual(sum(stage["days"] for stage in company["stages"]), company["estimated_days"])

    def test_location_and_project_conditions_change_explainable_price(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(
            template,
            title="ТЗ: Концепт геологии",
            input_data=DraftInputData(goal="Подготовить концепт геологии", source_data_ready=True),
        )
        document.requisites.update({"stages": ["Сбор данных", "Интерпретация", "Отчёт"], "city": "Москва"})
        office = procurement_service.estimate_for_tz(document)["estimate"]

        remote = document.model_copy(deep=True)
        remote.requisites["city"] = "Ямал, удалённое месторождение"
        remote.input_data.needs_3d_model = True
        field = procurement_service.estimate_for_tz(remote)["estimate"]

        self.assertIsNotNone(office)
        self.assertIsNotNone(field)
        office_prices = {row["company_id"]: row["cost_without_vat"] for row in office["companies"]}
        for company in field["companies"]:
            self.assertGreater(company["cost_without_vat"], office_prices[company["company_id"]])
            factor_keys = {factor["key"] for factor in company["cost_factors"]}
            self.assertIn("location", factor_keys)
            self.assertIn("3d_model", factor_keys)
            self.assertGreater(company["location_factor"], 1)
        self.assertEqual(field["summary"]["location"], "Ямал, удалённое месторождение")


class TZValidationTest(unittest.TestCase):
    def test_contractor_selection_completes_and_persists_tz(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(
            template,
            title="ТЗ: Концепт геологии",
            object_name="Северный участок",
            customer_name="Блок геологии",
            input_data=DraftInputData(
                goal="Подготовить геологическую концепцию",
                deadline="2029-12-31",
                source_data_ready=True,
            ),
        )
        document.requisites.update({
            "target_horizons": "Ю1–Ю4",
            "stages": list(template.stage_presets),
        })
        for section in document.sections:
            section.content = f"Заполненный раздел: {section.title}"

        estimate = procurement_service.estimate_for_tz(document)["estimate"]
        self.assertTrue(estimate["companies"])
        company = estimate["companies"][0]
        self.assertEqual(tz_validation_service.validate(document).ready_score, 100)

        with patch(
            "backend.app.api.routes.tz.tz_repository.update",
            new=AsyncMock(return_value=document),
        ) as update:
            response = asyncio.run(complete_document(TZCompleteRequest(
                document=document,
                company_id=company["company_id"],
            )))

        self.assertEqual(response.document.status, "ready")
        self.assertEqual(response.document.ready_score, 100)
        self.assertEqual(response.document.executor_name, company["company_name"])
        self.assertEqual(
            response.document.requisites["selected_contractor_id"],
            company["company_id"],
        )
        update.assert_awaited_once()

    def test_validation_returns_actionable_3d_and_required_field_issues(self):
        template = tz_template_service.get_template("tz-ptd-reserves")
        document = tz_generator.new_document(
            template,
            input_data=DraftInputData(needs_3d_model=True),
        )
        result = tz_validation_service.validate(document)
        codes = {issue.code for issue in result.issues}
        self.assertFalse(result.valid)
        self.assertIn("missing_object_name", codes)
        self.assertIn("3d_missing_input_data", codes)
        self.assertTrue(all(issue.recommendation for issue in result.issues))

    def test_template_specific_required_field_is_validated_and_used(self):
        template = tz_template_service.get_template("tz-geology-concept")
        document = tz_generator.new_document(
            template,
            object_name="Северный участок",
            customer_name="Блок геологии",
            input_data=DraftInputData(goal="Уточнить строение", deadline="2026-12-20"),
        )
        result = tz_validation_service.validate(document)
        self.assertIn("missing_template_field_target_horizons", {i.code for i in result.issues})
        document.requisites["target_horizons"] = "Юрские пласты Ю1–Ю4"
        tz_generator.generate(document, mode="full", template=template)
        scope = next(section.content for section in document.sections if section.key == "scope")
        self.assertIn("Юрские пласты Ю1–Ю4", scope)


if __name__ == "__main__":
    unittest.main()
