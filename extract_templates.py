# -*- coding: utf-8 -*-
import sys, io, os, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

base = r"Файлы\Выгрузка из системы\Шаблоны ТЗ"
files = []
for root, dirs, fs in os.walk(base):
    for f in fs:
        if f.lower().endswith('.docx'):
            files.append(os.path.join(root, f))
files.append(r"Файлы\Приложение № 2.1 Форма Технического задания.docx") if False else None

for path in files:
    print("="*90)
    print("FILE:", os.path.basename(path))
    print("="*90)
    try:
        doc = Document(path)
    except Exception as e:
        print("  <error>", e); continue
    # print paragraphs that look like headings (style name contains Heading OR bold/numbered)
    count=0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name if p.style else "") or ""
        is_head = "head" in style.lower() or "заголов" in style.lower()
        # heuristics: short lines with number prefix
        if is_head or (len(t) < 90 and (t[:2].strip().rstrip('.').isdigit() or t.isupper())):
            print(f"  [{style}] {t[:110]}")
            count+=1
        if count>60:
            break
    print(f"  (paragraphs total: {len(doc.paragraphs)}, tables: {len(doc.tables)})")
